"""DOCX parsing.

Easier than PDF: the format actually records structure. Headings carry a style
name, and tables are real table objects rather than aligned whitespace, so
table detection is exact here rather than heuristic.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.knowledge.parsers.base import Block, BlockType, ParsedDocument, ParseError


class DocxParser:
    @property
    def suffixes(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, path: Path) -> ParsedDocument:
        try:
            document = docx.Document(str(path))
        except (OSError, ValueError, KeyError) as exc:
            # python-docx raises KeyError on a zip that is not a .docx.
            raise ParseError(f"could not read DOCX {path.name}") from exc

        blocks: list[Block] = []
        heading: str | None = None

        for item in _in_document_order(document):
            if isinstance(item, Table):
                text = _table_text(item)
                if text:
                    blocks.append(Block(type=BlockType.TABLE, text=text, heading=heading))
                continue

            text = item.text.strip()
            if not text:
                continue

            style = (item.style.name or "") if item.style else ""
            if style.startswith("Heading"):
                heading = text
                blocks.append(
                    Block(
                        type=BlockType.HEADING,
                        text=text,
                        level=_heading_level(style),
                        heading=heading,
                    )
                )
            elif style.startswith("List"):
                blocks.append(Block(type=BlockType.LIST, text=text, heading=heading))
            else:
                blocks.append(Block(type=BlockType.PARAGRAPH, text=text, heading=heading))

        title = next((b.text for b in blocks if b.type is BlockType.HEADING), None)
        return ParsedDocument(blocks=blocks, title=title or path.stem)


def _in_document_order(document: DocxDocument) -> list[Paragraph | Table]:
    """Paragraphs and tables interleaved as they appear.

    python-docx exposes `.paragraphs` and `.tables` as separate lists, which loses
    their relative order — a table would end up detached from the text introducing
    it. Walking the body XML preserves it.
    """
    items: list[Paragraph | Table] = []
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            items.append(Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            items.append(Table(child, document))
    return items


def _table_text(table: Table) -> str:
    """Render a table as pipe-delimited rows, header row first.

    The rendering keeps column alignment readable to a model, which matters
    because the whole table becomes one chunk.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _heading_level(style_name: str) -> int | None:
    """ "Heading 2" -> 2."""
    tail = style_name.replace("Heading", "").strip()
    return int(tail) if tail.isdigit() else None
