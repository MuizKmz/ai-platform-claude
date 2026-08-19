"""Generate test documents.

Fixtures are generated rather than committed as binaries so they can be inspected,
adjusted, and diffed. Committing a 200-page PDF would put an opaque blob in git
that nobody can review, and the generator states the intent of each fixture in a
way the binary cannot.

    uv run python tests/fixtures/make_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

OUT = Path(__file__).parent / "documents"

# The table every table test asserts on. Split it anywhere and the rows below the
# break become numbers with no column names — the failure the chunker must avoid.
PART_TABLE = [
    ["Part Number", "Description", "Unit Price", "Lead Time"],
    ["XR-7742-B", "Hydraulic seal kit", "$284.50", "14 days"],
    ["QN-1183-A", "Thermal sensor array", "$1,290.00", "28 days"],
    ["ZK-9021-C", "Control board revision C", "$3,410.75", "45 days"],
    ["MB-4460-D", "Mounting bracket, steel", "$62.20", "7 days"],
]

PAGE_HEADER = "ACME INDUSTRIAL SUPPLIES - CONFIDENTIAL"


def _markdown_table(rows: list[list[str]]) -> str:
    header, *body = rows
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    lines += ["| " + " | ".join(row) + " |" for row in body]
    return "\n".join(lines)


def write_markdown() -> Path:
    path = OUT / "catalogue.md"
    path.write_text(
        "# Parts Catalogue\n\n"
        "The following components are available for order. Prices exclude tax.\n\n"
        + _markdown_table(PART_TABLE)
        + "\n\n# Ordering\n\n"
        "Orders are placed through the procurement portal. Lead times begin on the\n"
        "day payment clears, not the day the order is submitted.\n",
        encoding="utf-8",
    )
    return path


def write_html() -> Path:
    """Includes nav/header/footer chrome, which the parser must discard."""
    path = OUT / "catalogue.html"
    rows = "\n".join(
        "      <tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>"
        for row in PART_TABLE[1:]
    )
    headers = "".join(f"<th>{cell}</th>" for cell in PART_TABLE[0])
    path.write_text(
        f"""<!doctype html>
<html>
<head><title>Parts Catalogue</title></head>
<body>
  <nav>Home | Products | Contact</nav>
  <header>Acme Industrial Supplies</header>
  <main>
    <h1>Parts Catalogue</h1>
    <p>The following components are available for order. Prices exclude tax.</p>
    <table>
      <tr>{headers}</tr>
{rows}
    </table>
    <h2>Ordering</h2>
    <p>Orders are placed through the procurement portal.</p>
    <ul><li>Payment clears first</li><li>Then lead time begins</li></ul>
  </main>
  <footer>Copyright 2026 Acme Industrial Supplies</footer>
</body>
</html>
""",
        encoding="utf-8",
    )
    return path


def write_docx() -> Path:
    import docx

    path = OUT / "catalogue.docx"
    document = docx.Document()
    document.add_heading("Parts Catalogue", level=1)
    document.add_paragraph("The following components are available for order.")

    table = document.add_table(rows=len(PART_TABLE), cols=len(PART_TABLE[0]))
    for row_index, row in enumerate(PART_TABLE):
        for cell_index, value in enumerate(row):
            table.cell(row_index, cell_index).text = value

    document.add_heading("Ordering", level=1)
    document.add_paragraph("Orders are placed through the procurement portal.")
    document.add_paragraph("Payment clears first", style="List Bullet")
    document.save(str(path))
    return path


def write_pdf(pages: int = 200) -> Path:
    """A 200-page PDF with a repeated header, page numbers, and one table.

    Three properties are deliberate, and each maps to a test:
      - the header repeats on every page, so it must be recognised as chrome
      - page numbers differ per page, so they need separate handling
      - page 3 holds a column-aligned table that must survive chunking whole
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    path = OUT / "handbook.pdf"
    pdf = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER

    for page_number in range(1, pages + 1):
        pdf.setFont("Helvetica", 10)
        pdf.drawString(56, height - 50, PAGE_HEADER)

        y = height - 90
        if page_number == 3:
            pdf.drawString(56, y, "Parts Catalogue")
            y -= 28
            # Monospaced, so the column-gap heuristic sees aligned columns.
            pdf.setFont("Courier", 9)
            for row in PART_TABLE:
                pdf.drawString(56, y, f"{row[0]:<14}{row[1]:<28}{row[2]:<12}{row[3]}")
                y -= 14
            pdf.setFont("Helvetica", 10)
        else:
            pdf.drawString(56, y, f"Section {page_number}")
            y -= 28
            pdf.drawString(56, y, f"This is the body text of page {page_number}.")
            y -= 14
            pdf.drawString(56, y, "It describes routine operating procedure.")

        pdf.drawString(width / 2, 40, str(page_number))
        pdf.showPage()

    pdf.save()
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for writer in (write_markdown, write_html, write_docx, write_pdf):
        path = writer()
        print(f"wrote {path.name} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
